"""
M02 syllabus export task — PDF / DOCX / JSON.

Runs on the celery-heavy queue.  Export is only permitted for
FACULTY_APPROVED and ADMIN_LOCKED syllabi; the service layer enforces
this before dispatching, and this task double-checks on entry.

Data loaded per export:
  - Syllabus + outcomes (with CO-PO mappings) + units + references
  - Course metadata from M01 (code, title, credits, semester, L/T/P)
  - Programme outcomes from M01 (for CO-PO matrix column headers)
  - Only is_confirmed=True references appear in exports
"""
from __future__ import annotations

import asyncio
import io
import json
import logging
import re
import sys
import uuid as uuid_module
from datetime import timezone
from uuid import UUID

from app.database import tenant_schema_scope
from app.workers.base_task import VidyaTask
from app.workers.celery_app import celery_app

logger = logging.getLogger("vidya.worker.m02.export")

_async_engine = None


def _get_async_engine():
    global _async_engine
    if _async_engine is None:
        from sqlalchemy.ext.asyncio import create_async_engine
        from app.config import settings
        _async_engine = create_async_engine(settings.DATABASE_URL, pool_pre_ping=True)
        # Re-apply the schema at the START OF EVERY TRANSACTION, not once per
        # session. A commit hands this connection back — NullPool closes it, a pool
        # recycles it — so anything after the first commit would otherwise run with
        # search_path = public, and a pooled connection could arrive still carrying
        # ANOTHER tenant's search_path. A commit cannot undo a per-BEGIN SET LOCAL.
        from app.database import bind_tenant_search_path
        bind_tenant_search_path(_async_engine)
    return _async_engine


# ---------------------------------------------------------------------------
# Celery task
# ---------------------------------------------------------------------------

@celery_app.task(
    base=VidyaTask,
    name="app.workers.heavy.export_syllabus",
    autoretry_for=(ConnectionError, TimeoutError, OSError),
    max_retries=3,
    retry_backoff=True,
    retry_backoff_max=600,
    retry_jitter=True,
)
def export_syllabus(
    *,
    job_id: str,
    syllabus_id: str,
    tenant_id: str,
    schema_name: str,
    export_format: str,
    requested_by_user_id: str,
    request_id: str | None = None,
    **kwargs,
) -> dict:
    if sys.platform == "win32":
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    # Which tenant every transaction in this task belongs to. Held for the whole
    # run and dropped at the end of it: a worker process is long-lived and serves
    # every tenant in turn, and a schema left set is one the next task inherits.
    with tenant_schema_scope(schema_name):
        return asyncio.run(
            _run_export(
                syllabus_id=UUID(syllabus_id),
                tenant_id=UUID(tenant_id),
                schema_name=schema_name,
                export_format=export_format,
                requested_by_user_id=UUID(requested_by_user_id),
            )
        )


# ---------------------------------------------------------------------------
# Inner async implementation
# ---------------------------------------------------------------------------

async def _run_export(
    syllabus_id: UUID,
    tenant_id: UUID,
    schema_name: str,
    export_format: str,
    requested_by_user_id: UUID,
) -> dict:
    from sqlalchemy import select
    from sqlalchemy.ext.asyncio import AsyncSession

    from app.core.audit_log.models import AuditEventType
    from app.core.audit_log.service import AuditService
    from app.core.storage.models import StorageEntityType
    from app.core.storage.repository import StorageRepository
    from app.modules.m01_program_advisor.models import ProgramOutcome
    from app.modules.m01_program_advisor.repository import CourseRepository
    from app.modules.m02_syllabus.models import SyllabusStatus
    from app.modules.m02_syllabus.repository import SyllabusRepository

    tenant_slug = schema_name.removeprefix("tenant_")
    engine = _get_async_engine()

    _EXPORTABLE = {SyllabusStatus.APPROVED, SyllabusStatus.LOCKED}

    try:
        async with AsyncSession(engine, expire_on_commit=False) as session:
            syllabus = await SyllabusRepository.get_detail(syllabus_id, db=session)
            if syllabus is None:
                raise ValueError(f"Syllabus {syllabus_id} not found.")
            if syllabus.status not in _EXPORTABLE:
                raise ValueError(
                    f"Syllabus {syllabus_id} is {syllabus.status.value}; "
                    "export requires an APPROVED or LOCKED official syllabus."
                )

            course = await CourseRepository.get_by_id(syllabus.course_id, db=session)
            if course is None:
                raise ValueError(f"Course {syllabus.course_id} not found in M01.")

            # Load programme outcomes for CO-PO matrix column headers
            po_stmt = (
                select(ProgramOutcome)
                .where(ProgramOutcome.program_id == course.program_id)
                .order_by(ProgramOutcome.display_order)
            )
            pos = list((await session.execute(po_stmt)).scalars().all())

            buf = io.BytesIO()
            if export_format == "pdf":
                content_type = "application/pdf"
                ext = "pdf"
                _generate_pdf(buf, syllabus, course, pos)
            elif export_format == "docx":
                content_type = (
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                )
                ext = "docx"
                _generate_docx(buf, syllabus, course, pos)
            else:
                content_type = "application/json"
                ext = "json"
                _generate_json(buf, syllabus, course, pos)

            buf.seek(0)
            file_bytes = buf.read()
            size_bytes = len(file_bytes)

            safe_code  = re.sub(r"[^a-z0-9_-]", "_", course.code.lower())[:20].strip("_")
            filename   = f"syllabus_{safe_code}_v{syllabus.version}.{ext}"
            file_uuid  = uuid_module.uuid4()
            object_key = (
                f"vidya-assets/{tenant_slug}/syllabus_export"
                f"/{syllabus_id}/{file_uuid}-{filename}"
            )

            await asyncio.to_thread(_s3_put_object, object_key, file_bytes, content_type)

            asset = await StorageRepository.create(
                uploaded_by_user_id=requested_by_user_id,
                entity_type=StorageEntityType.SYLLABUS_EXPORT.value,
                entity_id=syllabus_id,
                object_key=object_key,
                original_filename=filename,
                size_bytes=size_bytes,
                content_type=content_type,
                db=session,
            )
            await session.commit()

        download_url = await StorageRepository.generate_presigned_get_url(
            object_key=object_key,
            expires_in_seconds=86_400,
        )

        await AuditService.log(
            AuditEventType.SYLLABUS_EXPORT_COMPLETED,
            actor_user_id=requested_by_user_id,
            actor_role="SYSTEM",
            tenant_id=tenant_id,
            schema_name=schema_name,
            target_entity="Syllabus",
            target_id=str(syllabus_id),
            metadata={
                "asset_id":   str(asset.id),
                "format":     export_format,
                "size_bytes": size_bytes,
            },
        )

        logger.info(
            "m02.export: %s export complete (syllabus=%s size=%d)",
            export_format, syllabus_id, size_bytes,
        )

        return {
            "download_url": download_url,
            "asset_id":     str(asset.id),
            "format":       export_format,
            "size_bytes":   size_bytes,
        }

    except Exception as exc:
        try:
            await AuditService.log(
                AuditEventType.SYLLABUS_EXPORT_FAILED,
                actor_user_id=requested_by_user_id,
                actor_role="SYSTEM",
                tenant_id=tenant_id,
                schema_name=schema_name,
                target_entity="Syllabus",
                target_id=str(syllabus_id),
                metadata={"error": str(exc)[:500], "format": export_format},
            )
        except Exception:
            logger.exception("m02.export: failed to log SYLLABUS_EXPORT_FAILED audit")
        raise


# ---------------------------------------------------------------------------
# S3 direct upload
# ---------------------------------------------------------------------------

def _s3_put_object(object_key: str, file_bytes: bytes, content_type: str) -> None:
    import boto3
    from app.config import settings

    client = boto3.client(
        "s3",
        endpoint_url=settings.S3_ENDPOINT,
        aws_access_key_id=settings.S3_ACCESS_KEY,
        aws_secret_access_key=settings.S3_SECRET_KEY,
        region_name=settings.S3_REGION,
        use_ssl=settings.S3_USE_SSL,
    )
    s3_key = object_key.removeprefix(f"{settings.S3_BUCKET}/")
    client.put_object(
        Bucket=settings.S3_BUCKET,
        Key=s3_key,
        Body=file_bytes,
        ContentType=content_type,
    )


# ---------------------------------------------------------------------------
# PDF generation (reportlab) — the OFFICIAL university regulation document
#
# This is not a summary of a syllabus. It is the syllabus: the page that goes into
# the university's regulation handbook, is approved by the Board of Studies, and
# is handed to faculty and students as the definitive statement of what the course
# teaches. It has to look like one.
#
# Layout, in the order a regulation prints:
#
#     COURSE CODE - COURSE NAME
#     Course Information table (Code / Name / Credits / L-T-P / Contact Hours / Category)
#     COURSE OBJECTIVES
#     COURSE OUTCOMES              CO1..CO5 with Bloom levels
#     CO-PO MAPPING MATRIX
#     UNIT I .. UNIT V             each a prose block with its hour allocation
#     PRACTICAL COMPONENTS         (only if the course has practical hours)
#     INTERNAL ASSESSMENT          (only if suggested)
#     TEXT BOOKS / REFERENCE BOOKS / SUGGESTED READING / WEB RESOURCES
# ---------------------------------------------------------------------------

def _generate_pdf(buf, syllabus, course, pos):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    from app.modules.m02_syllabus.formatting import (
        course_information,
        group_references,
        roman,
        unit_topic_lines,
    )

    _INK      = colors.HexColor("#111111")
    _RULE     = colors.HexColor("#333333")
    _BAND     = colors.HexColor("#E8E8E8")
    _STRENGTH = {"HIGH": "H", "MEDIUM": "M", "LOW": "L"}

    styles = getSampleStyleSheet()

    doc_title = ParagraphStyle(
        "docTitle", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=13, leading=16, alignment=1, textColor=_INK,
    )
    section = ParagraphStyle(
        "section", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=10, leading=13, spaceBefore=10, spaceAfter=5, textColor=_INK,
    )
    unit_head = ParagraphStyle(
        "unitHead", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=9.5, leading=12, spaceBefore=8, spaceAfter=3, textColor=_INK,
    )
    # Justified body text is what makes a page read as a regulation rather than as
    # a printout: the unit blocks are prose, and prose in an official document is
    # set flush on both edges.
    body = ParagraphStyle(
        "body", parent=styles["Normal"], fontName="Helvetica",
        fontSize=9, leading=13, alignment=TA_JUSTIFY, textColor=_INK,
    )
    listed = ParagraphStyle(
        "listed", parent=body, leftIndent=14, spaceAfter=2, alignment=0,
    )
    # One syllabus line. Tight leading, hanging indent — a regulation packs 12-20 of
    # these under each unit heading and they must not sprawl over the page.
    topic = ParagraphStyle(
        "topic", parent=styles["Normal"], fontName="Helvetica", fontSize=9,
        leading=11.5, leftIndent=16, firstLineIndent=-10, spaceAfter=1, textColor=_INK,
    )
    tiny = ParagraphStyle("tiny", parent=styles["Normal"], fontSize=7, leading=9)

    grid = TableStyle([
        ("FONTNAME",      (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE",      (0, 0), (-1, -1), 8.5),
        ("BOX",           (0, 0), (-1, -1), 0.6, _RULE),
        ("INNERGRID",     (0, 0), (-1, -1), 0.3, _RULE),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING",   (0, 0), (-1, -1), 5),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ])

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=1.8 * cm, bottomMargin=1.8 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
        title=f"{course.code} {course.title}",
    )
    story = []

    def rule(thickness=1.0):
        story.append(HRFlowable(width="100%", thickness=thickness, color=_RULE))

    # ── Title ────────────────────────────────────────────────────────────────
    story.append(Paragraph(f"{course.code} &nbsp;-&nbsp; {course.title.upper()}", doc_title))
    story.append(Spacer(1, 0.15 * cm))
    rule(1.2)
    story.append(Spacer(1, 0.3 * cm))

    # ── Course Information ───────────────────────────────────────────────────
    # Every value here is DERIVED from the course row — the printed page cannot
    # disagree with the curriculum it belongs to. Contact Hours is the one figure the
    # Board may state for itself, and when it has, the page prints what it said.
    info = course_information(
        course,
        teaching_hours=syllabus.teaching_hours,
        hours_per_week=syllabus.hours_per_week,
    )
    info_rows = [
        ["Course Code",   info["course_code"],   "Credits",              str(info["credits"])],
        ["Course Name",   info["course_name"],   "L-T-P",                info["ltp"]],
        ["Category",      info["category"],      "Total Teaching Hours", str(info["contact_hours"])],
        ["",              "",                    "Hours / Week",         str(info["hours_per_week"])],
    ]
    t = Table(info_rows, colWidths=[3 * cm, 6.4 * cm, 3 * cm, 4.6 * cm])
    t.setStyle(TableStyle(grid.getCommands() + [
        ("BACKGROUND", (0, 0), (0, -1), _BAND),
        ("BACKGROUND", (2, 0), (2, -1), _BAND),
        ("FONTNAME",   (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME",   (2, 0), (2, -1), "Helvetica-Bold"),
    ]))
    story.append(t)

    # ── Course Objectives ────────────────────────────────────────────────────
    objectives = list(syllabus.objectives or [])
    if objectives:
        story.append(Paragraph("COURSE OBJECTIVES", section))
        for i, obj in enumerate(objectives, 1):
            story.append(Paragraph(f"{i}. {obj}", listed))

    # ── Course Outcomes ──────────────────────────────────────────────────────
    cos_sorted = sorted(syllabus.outcomes, key=lambda c: c.display_order)
    if cos_sorted:
        story.append(Paragraph("COURSE OUTCOMES", section))
        story.append(Paragraph(
            "On successful completion of this course, the student will be able to:", body,
        ))
        story.append(Spacer(1, 0.15 * cm))
        rows = [["CO", "Course Outcome", "Bloom's Level"]] + [
            [
                co.code,
                Paragraph(co.description, body),
                co.bloom_level.value.title() if co.bloom_level else "-",
            ]
            for co in cos_sorted
        ]
        t = Table(rows, colWidths=[1.6 * cm, 12.4 * cm, 3 * cm])
        t.setStyle(TableStyle(grid.getCommands() + [
            ("BACKGROUND", (0, 0), (-1, 0), _BAND),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME",   (0, 1), (0, -1), "Helvetica-Bold"),
            ("ALIGN",      (0, 0), (0, -1), "CENTER"),
            ("ALIGN",      (2, 0), (2, -1), "CENTER"),
        ]))
        story.append(t)

    # ── CO-PO Mapping ────────────────────────────────────────────────────────
    if cos_sorted and pos:
        story.append(Paragraph("CO-PO MAPPING", section))
        index = {
            (co.id, m.po_id): m.mapping_strength.value
            for co in cos_sorted for m in co.mappings
        }
        header = ["CO"] + [p.code for p in pos]
        rows = [header] + [
            [co.code] + [_STRENGTH.get(index.get((co.id, p.id), ""), "-") for p in pos]
            for co in cos_sorted
        ]
        col = max(0.9 * cm, min(1.6 * cm, 15 * cm / (len(pos) + 1)))
        t = Table(rows, colWidths=[1.6 * cm] + [col] * len(pos))
        t.setStyle(TableStyle(grid.getCommands() + [
            ("BACKGROUND", (0, 0), (-1, 0), _BAND),
            ("BACKGROUND", (0, 0), (0, -1), _BAND),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME",   (0, 0), (0, -1), "Helvetica-Bold"),
            ("ALIGN",      (0, 0), (-1, -1), "CENTER"),
            ("FONTSIZE",   (0, 0), (-1, -1), 8),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.1 * cm))
        story.append(Paragraph("H = High   M = Medium   L = Low", tiny))

    # ── Units ────────────────────────────────────────────────────────────────
    # The heart of the document. Each unit is a heading with its hour allocation,
    # then a single justified prose block naming the concepts it teaches.
    units_sorted = sorted(syllabus.units, key=lambda u: u.unit_number)
    if units_sorted:
        story.append(Spacer(1, 0.25 * cm))
        rule(0.8)
        total = sum(u.total_hours or 0 for u in units_sorted)

        for unit in units_sorted:
            hours = f"{unit.total_hours} Hours" if unit.total_hours else ""
            head = Table(
                [[
                    Paragraph(
                        f"UNIT {roman(unit.unit_number)} &nbsp;-&nbsp; {(unit.title or '').upper()}",
                        unit_head,
                    ),
                    Paragraph(hours, ParagraphStyle(
                        "hrs", parent=unit_head, alignment=2,
                    )),
                ]],
                colWidths=[13.5 * cm, 3.5 * cm],
            )
            head.setStyle(TableStyle([
                ("LEFTPADDING",  (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING",   (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 2),
                ("LINEBELOW",    (0, 0), (-1, -1), 0.4, _RULE),
                ("VALIGN",       (0, 0), (-1, -1), "BOTTOM"),
            ]))
            story.append(head)

            # The unit's academic topics, one per line — 12 to 20 of them in a real
            # regulation. This IS the unit: it is what a lecturer reads to know what
            # to teach, and what a student reads to know what they will be taught.
            lines = unit_topic_lines(unit)
            if lines:
                story.append(Spacer(1, 0.12 * cm))
                for line in lines:
                    story.append(Paragraph(f"&bull;&nbsp;&nbsp;{line}", topic))

        if total:
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph(f"<b>TOTAL: {total} HOURS</b>", body))
        rule(0.8)

    # ── Practical Components ─────────────────────────────────────────────────
    practicals = list(syllabus.practical_components or [])
    if practicals:
        story.append(Paragraph("PRACTICAL COMPONENTS", section))
        for i, item in enumerate(practicals, 1):
            story.append(Paragraph(f"{i}. {item}", listed))

    # ── Internal Assessment ──────────────────────────────────────────────────
    assessment = list(getattr(syllabus, "internal_assessment", None) or [])
    if assessment:
        story.append(Paragraph("INTERNAL ASSESSMENT", section))
        for item in assessment:
            story.append(Paragraph(f"&bull; {item}", listed))

    # ── Bibliography — four sections, empty ones omitted ──────────────────────
    # Only CONFIRMED references print: the AI never invents bibliographic detail,
    # so an unconfirmed row is a search result nobody has vouched for and has no
    # place in a published regulation.
    confirmed = [r for r in syllabus.references if r.is_confirmed]
    for heading, refs in group_references(confirmed).items():
        story.append(Paragraph(heading.upper(), section))
        for i, ref in enumerate(refs, 1):
            authors = ", ".join(ref.authors) if ref.authors else ""
            year    = f"({ref.year})" if ref.year else ""
            line    = ". ".join(x for x in [f"{i}. {authors}".rstrip(". "), f"{year} {ref.title}".strip()] if x.strip())
            if ref.publisher:
                line += f", {ref.publisher}"
            if ref.isbn:
                line += f". ISBN: {ref.isbn}"
            elif ref.doi:
                line += f". DOI: {ref.doi}"
            elif ref.url:
                line += f". {ref.url}"
            story.append(Paragraph(line + ".", listed))

    doc.build(story)


# ---------------------------------------------------------------------------
# DOCX generation (python-docx) — the same official regulation document
#
# Deliberately mirrors the PDF section for section. A Board that exports one
# format and then the other should get the same document, not two different
# takes on it.
# ---------------------------------------------------------------------------

def _generate_docx(buf, syllabus, course, pos):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    from app.modules.m02_syllabus.formatting import (
        course_information,
        group_references,
        roman,
        unit_topic_lines,
    )

    _STRENGTH = {"HIGH": "H", "MEDIUM": "M", "LOW": "L"}

    doc = Document()

    def heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        return p

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{course.code} - {course.title.upper()}")
    run.bold = True
    run.font.size = Pt(14)

    # ── Course Information ───────────────────────────────────────────────────
    info = course_information(
        course,
        teaching_hours=syllabus.teaching_hours,
        hours_per_week=syllabus.hours_per_week,
    )
    t = doc.add_table(rows=0, cols=4)
    t.style = "Table Grid"
    for left, lval, right, rval in [
        ("Course Code", info["course_code"], "Credits",              str(info["credits"])),
        ("Course Name", info["course_name"], "L-T-P",                info["ltp"]),
        ("Category",    info["category"],    "Total Teaching Hours", str(info["contact_hours"])),
        ("",            "",                  "Hours / Week",         str(info["hours_per_week"])),
    ]:
        row = t.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = left, lval, right, rval
        for idx in (0, 2):
            for para in row[idx].paragraphs:
                for r in para.runs:
                    r.bold = True

    # ── Course Objectives ────────────────────────────────────────────────────
    objectives = list(syllabus.objectives or [])
    if objectives:
        heading("Course Objectives")
        for obj in objectives:
            doc.add_paragraph(obj, style="List Number")

    # ── Course Outcomes ──────────────────────────────────────────────────────
    cos_sorted = sorted(syllabus.outcomes, key=lambda c: c.display_order)
    if cos_sorted:
        heading("Course Outcomes")
        doc.add_paragraph(
            "On successful completion of this course, the student will be able to:"
        )
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "CO", "Course Outcome", "Bloom's Level"
        for cell in hdr:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.bold = True
        for co in cos_sorted:
            row = t.add_row().cells
            row[0].text = co.code
            row[1].text = co.description
            row[2].text = co.bloom_level.value.title() if co.bloom_level else "-"

    # ── CO-PO Mapping ────────────────────────────────────────────────────────
    if cos_sorted and pos:
        heading("CO-PO Mapping")
        index = {
            (co.id, m.po_id): m.mapping_strength.value
            for co in cos_sorted for m in co.mappings
        }
        t = doc.add_table(rows=1, cols=len(pos) + 1)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        hdr[0].text = "CO"
        for i, po in enumerate(pos, 1):
            hdr[i].text = po.code
        for co in cos_sorted:
            row = t.add_row().cells
            row[0].text = co.code
            for i, po in enumerate(pos, 1):
                row[i].text = _STRENGTH.get(index.get((co.id, po.id), ""), "-")
        doc.add_paragraph("H = High   M = Medium   L = Low")

    # ── Units ────────────────────────────────────────────────────────────────
    units_sorted = sorted(syllabus.units, key=lambda u: u.unit_number)
    total = sum(u.total_hours or 0 for u in units_sorted)
    for unit in units_sorted:
        hours = f"   ({unit.total_hours} Hours)" if unit.total_hours else ""
        heading(f"Unit {roman(unit.unit_number)} - {unit.title}{hours}")
        for line in unit_topic_lines(unit):
            doc.add_paragraph(line, style="List Bullet")
    if total:
        p = doc.add_paragraph()
        p.add_run(f"TOTAL: {total} HOURS").bold = True

    # ── Practical Components ─────────────────────────────────────────────────
    practicals = list(syllabus.practical_components or [])
    if practicals:
        heading("Practical Components")
        for item in practicals:
            doc.add_paragraph(item, style="List Number")

    # ── Internal Assessment ──────────────────────────────────────────────────
    assessment = list(getattr(syllabus, "internal_assessment", None) or [])
    if assessment:
        heading("Internal Assessment")
        for item in assessment:
            doc.add_paragraph(item, style="List Bullet")

    # ── Bibliography — four sections, empty ones omitted ──────────────────────
    confirmed = [r for r in syllabus.references if r.is_confirmed]
    for section_name, refs in group_references(confirmed).items():
        heading(section_name)
        for ref in refs:
            authors = ", ".join(ref.authors) if ref.authors else ""
            year    = f"({ref.year}) " if ref.year else ""
            line    = f"{authors}. {year}{ref.title}".strip(". ")
            if ref.publisher:
                line += f", {ref.publisher}"
            if ref.isbn:
                line += f". ISBN: {ref.isbn}"
            elif ref.doi:
                line += f". DOI: {ref.doi}"
            elif ref.url:
                line += f". {ref.url}"
            doc.add_paragraph(line + ".", style="List Number")

    doc.save(buf)


# ---------------------------------------------------------------------------
# JSON generation (machine-readable structured export)
# ---------------------------------------------------------------------------

def _generate_json(buf, syllabus, course, pos):
    from datetime import datetime

    def _dt(d):
        return d.isoformat() if d else None

    mapping_index: dict[tuple, dict] = {}
    for co in syllabus.outcomes:
        for m in co.mappings:
            mapping_index[(co.id, m.po_id)] = {
                "strength":     m.mapping_strength.value,
                "justification": m.justification,
            }

    cos_sorted   = sorted(syllabus.outcomes, key=lambda c: c.display_order)
    units_sorted = sorted(syllabus.units,    key=lambda u: u.unit_number)

    document = {
        "export_metadata": {
            "exported_at":  datetime.now(timezone.utc).isoformat(),
            "syllabus_id":  str(syllabus.id),
            "course_id":    str(syllabus.course_id),
            "version":      syllabus.version,
            "status":       syllabus.status.value,
            "approved_at":  _dt(syllabus.approved_at),
            "locked_at":    _dt(syllabus.locked_at),
            "ai_model":     syllabus.ai_model,
        },
        "course": {
            "id":              str(course.id),
            "code":            course.code,
            "title":           course.title,
            "credits":         course.credits,
            "semester":        course.semester,
            "hours_lecture":   course.hours_lecture,
            "hours_tutorial":  course.hours_tutorial,
            "hours_practical": course.hours_practical,
            "description":     course.description,
        },
        "course_outcomes": [
            {
                "id":            str(co.id),
                "code":          co.code,
                "description":   co.description,
                "bloom_level":   co.bloom_level.value if co.bloom_level else None,
                "display_order": co.display_order,
            }
            for co in cos_sorted
        ],
        "co_po_matrix": {
            "program_outcomes": [
                {
                    "id":            str(po.id),
                    "code":          po.code,
                    "description":   po.description,
                    "bloom_level":   po.bloom_level,
                    "display_order": po.display_order,
                }
                for po in pos
            ],
            "mappings": [
                {
                    "co_id":   str(co.id),
                    "co_code": co.code,
                    "po_mappings": {
                        str(po.id): mapping_index[(co.id, po.id)]
                        for po in pos
                        if (co.id, po.id) in mapping_index
                    },
                }
                for co in cos_sorted
            ],
        },
        # The official prose sections, so a JSON export is the same document as the
        # PDF rather than a lossy sibling of it.
        "objectives":           list(syllabus.objectives or []),
        "practical_components": list(syllabus.practical_components or []),
        "internal_assessment":  list(getattr(syllabus, "internal_assessment", None) or []),
        "units": [
            {
                "id":           str(unit.id),
                "unit_number":  unit.unit_number,
                "title":        unit.title,
                # `content` is what prints; `topics` is the structured scaffolding
                # underneath it that downstream generators read.
                "content":      unit.content,
                "total_hours":  unit.total_hours,
                "pedagogy":     unit.pedagogy,
                "topics":       unit.topics or [],
                "bloom_summary": unit.bloom_summary or {},
            }
            for unit in units_sorted
        ],
        "references": [
            {
                "id":           str(ref.id),
                "title":        ref.title,
                "authors":      ref.authors or [],
                "year":         ref.year,
                "ref_type":     ref.ref_type.value,
                "source":       ref.source.value if ref.source else None,
                "doi":          ref.doi,
                "isbn":         ref.isbn,
                "url":          ref.url,
                "publisher":    ref.publisher,
                "is_confirmed": ref.is_confirmed,
            }
            for ref in syllabus.references
            if ref.is_confirmed
        ],
    }

    buf.write(json.dumps(document, indent=2, ensure_ascii=False).encode("utf-8"))
