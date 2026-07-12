import asyncio
import io
import logging
import re
import sys
import uuid as uuid_module
from uuid import UUID

from app.workers.base_task import VidyaTask
from app.workers.celery_app import celery_app

logger = logging.getLogger("vidya.worker.m01.export")

_async_engine = None


def _get_async_engine():
    global _async_engine
    if _async_engine is None:
        from sqlalchemy.ext.asyncio import create_async_engine
        from sqlalchemy.pool import NullPool
        from app.config import settings
        # NullPool: no connection caching between asyncio.run() calls.
        # Prevents "Future attached to a different loop" on Windows --pool=solo.
        _async_engine = create_async_engine(settings.DATABASE_URL, poolclass=NullPool)
    return _async_engine


# ---------------------------------------------------------------------------
# Celery task
# ---------------------------------------------------------------------------

@celery_app.task(
    base=VidyaTask,
    name="app.workers.heavy.export_program",
    autoretry_for=(ConnectionError, TimeoutError, OSError),
    max_retries=3,
    retry_backoff=True,
    retry_backoff_max=600,
    retry_jitter=True,
)
def export_program(
    *,
    job_id: str,
    program_id: str,
    tenant_id: str,
    schema_name: str,
    export_format: str,
    requested_by_user_id: str,
    request_id: str | None = None,
    **kwargs,
) -> dict:
    if sys.platform == "win32":
        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
    return asyncio.run(
        _run_export(
            program_id=UUID(program_id),
            tenant_id=UUID(tenant_id),
            schema_name=schema_name,
            export_format=export_format,
            requested_by_user_id=UUID(requested_by_user_id),
        )
    )


# ---------------------------------------------------------------------------
# Inner async implementation
# ---------------------------------------------------------------------------

async def _run_export(
    program_id: UUID,
    tenant_id: UUID,
    schema_name: str,
    export_format: str,
    requested_by_user_id: UUID,
) -> dict:
    from sqlalchemy import text
    from sqlalchemy.ext.asyncio import AsyncSession

    from app.core.audit_log.models import AuditEventType
    from app.core.audit_log.service import AuditService
    from app.core.storage.models import StorageEntityType
    from app.core.storage.repository import StorageRepository
    from app.modules.m01_program_advisor.models import ProgramStatus
    from app.modules.m01_program_advisor.repository import (
        CoursePrerequisiteRepository,
        CourseRepository,
        ProgramOutcomeRepository,
        ProgramRepository,
    )

    tenant_slug = schema_name.removeprefix("tenant_")
    engine = _get_async_engine()

    try:
        async with AsyncSession(engine, expire_on_commit=False) as session:
            await session.execute(text(f"SET search_path TO {schema_name}, public"))

            program = await ProgramRepository.get_by_id(program_id, db=session)
            if program is None:
                raise ValueError(f"Program {program_id} not found.")
            # Exportable once the governance authority has locked it — and it stays
            # exportable after the Dean publishes it. PUBLISHED is the same frozen
            # content as APPROVED, one step later, so refusing to export it would
            # mean the final, live curriculum is the one document nobody can print.
            if program.status not in (ProgramStatus.APPROVED, ProgramStatus.PUBLISHED):
                raise ValueError(
                    f"Program {program_id} is {program.status.value}; only an "
                    f"approved (locked) or published curriculum can be exported."
                )

            outcomes = await ProgramOutcomeRepository.list_by_program(program_id, db=session)
            courses = await CourseRepository.list_by_program(program_id, db=session)

            course_code_map: dict[UUID, str] = {c.id: c.code for c in courses}
            prereqs_by_course: dict[UUID, list[str]] = {}
            for course in courses:
                rows = await CoursePrerequisiteRepository.list_by_course(course.id, db=session)
                prereqs_by_course[course.id] = [
                    course_code_map.get(r.prerequisite_course_id, str(r.prerequisite_course_id))
                    for r in rows
                ]

            buf = io.BytesIO()
            if export_format == "pdf":
                content_type = "application/pdf"
                ext = "pdf"
                _generate_pdf(buf, program, outcomes, courses, prereqs_by_course)
            else:
                content_type = (
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                )
                ext = "docx"
                _generate_docx(buf, program, outcomes, courses, prereqs_by_course)
            buf.seek(0)
            file_bytes = buf.read()
            size_bytes = len(file_bytes)

            safe_title = re.sub(r"[^a-z0-9_-]", "_", program.title.lower())[:40].strip("_")
            filename = f"{safe_title}_v{program.version}.{ext}"
            file_uuid = uuid_module.uuid4()
            object_key = (
                f"vidya-assets/{tenant_slug}/program_export"
                f"/{program_id}/{file_uuid}-{filename}"
            )

            await asyncio.to_thread(_s3_put_object, object_key, file_bytes, content_type)

            asset = await StorageRepository.create(
                uploaded_by_user_id=requested_by_user_id,
                entity_type=StorageEntityType.PROGRAM_EXPORT.value,
                entity_id=program_id,
                object_key=object_key,
                original_filename=filename,
                size_bytes=size_bytes,
                content_type=content_type,
                db=session,
            )
            await session.commit()

        download_url = await StorageRepository.generate_presigned_get_url(
            object_key=object_key,
            expires_in_seconds=86_400,
        )

        await AuditService.log(
            AuditEventType.PROGRAM_EXPORT_COMPLETED,
            actor_user_id=requested_by_user_id,
            actor_role="SYSTEM",
            tenant_id=tenant_id,
            schema_name=schema_name,
            target_entity="Program",
            target_id=str(program_id),
            metadata={
                "asset_id":   str(asset.id),
                "format":     export_format,
                "size_bytes": size_bytes,
            },
        )

        logger.info(
            "m01.export: %s export complete (program=%s size=%d)",
            export_format, program_id, size_bytes,
        )

        return {
            "download_url": download_url,
            "asset_id":     str(asset.id),
            "format":       export_format,
            "size_bytes":   size_bytes,
        }

    except Exception as exc:
        try:
            await AuditService.log(
                AuditEventType.PROGRAM_EXPORT_FAILED,
                actor_user_id=requested_by_user_id,
                actor_role="SYSTEM",
                tenant_id=tenant_id,
                schema_name=schema_name,
                target_entity="Program",
                target_id=str(program_id),
                metadata={"error": str(exc)[:500], "format": export_format},
            )
        except Exception:
            logger.exception("m01.export: failed to log PROGRAM_EXPORT_FAILED audit")
        raise


# ---------------------------------------------------------------------------
# S3 direct upload (worker has credentials — no presigned PUT needed)
# ---------------------------------------------------------------------------

def _s3_put_object(object_key: str, file_bytes: bytes, content_type: str) -> None:
    import boto3
    from app.config import settings

    client = boto3.client(
        "s3",
        endpoint_url=settings.S3_ENDPOINT,
        aws_access_key_id=settings.S3_ACCESS_KEY,
        aws_secret_access_key=settings.S3_SECRET_KEY,
        region_name=settings.S3_REGION,
        use_ssl=settings.S3_USE_SSL,
    )
    s3_key = object_key.removeprefix(f"{settings.S3_BUCKET}/")
    client.put_object(
        Bucket=settings.S3_BUCKET,
        Key=s3_key,
        Body=file_bytes,
        ContentType=content_type,
    )


# ---------------------------------------------------------------------------
# PDF generation (reportlab)
# ---------------------------------------------------------------------------

def _generate_pdf(buf, program, outcomes, courses, prereqs_by_course):
    from collections import defaultdict

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    _BLUE    = colors.HexColor("#4472C4")
    _ALT_ROW = colors.HexColor("#EBF0FA")
    _GRID_STYLE = [
        ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTNAME",     (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE",     (0, 0), (-1, -1), 8),
        ("BACKGROUND",   (0, 0), (-1, 0),  _BLUE),
        ("TEXTCOLOR",    (0, 0), (-1, 0),  colors.white),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, _ALT_ROW]),
        ("BOX",          (0, 0), (-1, -1), 0.5, colors.grey),
        ("INNERGRID",    (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING",   (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
    ]

    styles = getSampleStyleSheet()
    small  = styles["Normal"].clone("small")
    small.fontSize = 8
    small.leading  = 10

    doc   = SimpleDocTemplate(buf, pagesize=A4,
                              topMargin=2*cm, bottomMargin=2*cm,
                              leftMargin=2*cm, rightMargin=2*cm)
    story = []

    # -- Cover ---------------------------------------------------------------
    story.append(Paragraph(program.title, styles["Title"]))
    story.append(Spacer(1, 0.4*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=_BLUE))
    story.append(Spacer(1, 0.4*cm))

    meta_rows = [
        ["Degree Type",    program.degree_type],
        ["Department",     program.department],
        ["Duration",       f"{program.duration_years} year(s)"],
        ["Total Credits",  str(program.total_credits)],
        ["Version",        str(program.version)],
        ["Status",         program.status.value],
    ]
    if program.approved_at:
        meta_rows.append(["Approved At", program.approved_at.strftime("%Y-%m-%d")])

    meta_table = Table(meta_rows, colWidths=[4.5*cm, 10*cm])
    meta_table.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (0, -1), colors.lightgrey),
        ("FONTNAME",     (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE",     (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.whitesmoke, colors.white]),
        ("BOX",          (0, 0), (-1, -1), 0.5, colors.grey),
        ("INNERGRID",    (0, 0), (-1, -1), 0.25, colors.grey),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING",   (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 1*cm))

    # -- Program Outcomes ----------------------------------------------------
    if outcomes:
        story.append(Paragraph("Program Outcomes", styles["Heading1"]))
        story.append(Spacer(1, 0.3*cm))
        header = [["Code", "Description", "Bloom Level", "Order"]]
        rows = [
            [
                o.code,
                Paragraph(o.description, small),
                o.bloom_level or "—",
                str(o.display_order),
            ]
            for o in outcomes
        ]
        t = Table(header + rows, colWidths=[2*cm, 9.5*cm, 2.8*cm, 1.5*cm])
        t.setStyle(TableStyle(_GRID_STYLE))
        story.append(t)
        story.append(Spacer(1, 1*cm))

    # -- Course Structure (semester-grouped) ---------------------------------
    story.append(Paragraph("Course Structure", styles["Heading1"]))
    story.append(Spacer(1, 0.3*cm))

    by_semester = defaultdict(list)
    for c in courses:
        by_semester[c.semester].append(c)

    for sem in sorted(by_semester.keys()):
        story.append(Paragraph(f"Semester {sem}", styles["Heading2"]))
        header = [["Code", "Title", "Cr", "Type", "L/T/P", "Prerequisites"]]
        rows = []
        for c in by_semester[sem]:
            ltp      = f"{c.hours_lecture or 0}/{c.hours_tutorial or 0}/{c.hours_practical or 0}"
            prereqs  = ", ".join(prereqs_by_course.get(c.id, [])) or "—"
            ctype    = "Elective" if c.is_elective else "Core"
            rows.append([
                c.code,
                Paragraph(c.title, small),
                str(c.credits),
                ctype,
                ltp,
                Paragraph(prereqs, small),
            ])
        t = Table(
            header + rows,
            colWidths=[2.2*cm, 5.2*cm, 1*cm, 1.8*cm, 1.6*cm, 4*cm],
        )
        t.setStyle(TableStyle(_GRID_STYLE))
        story.append(t)
        story.append(Spacer(1, 0.5*cm))

    doc.build(story)


# ---------------------------------------------------------------------------
# DOCX generation (python-docx)
# ---------------------------------------------------------------------------

def _generate_docx(buf, program, outcomes, courses, prereqs_by_course):
    from collections import defaultdict

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # -- Cover ---------------------------------------------------------------
    title_para = doc.add_heading(program.title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Program Details", 1)
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = "Table Grid"
    meta_rows = [
        ("Degree Type",   program.degree_type),
        ("Department",    program.department),
        ("Duration",      f"{program.duration_years} year(s)"),
        ("Total Credits", str(program.total_credits)),
        ("Version",       str(program.version)),
        ("Status",        program.status.value),
        ("Approved At",   program.approved_at.strftime("%Y-%m-%d") if program.approved_at else "—"),
    ]
    for label, value in meta_rows:
        row = meta_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    doc.add_paragraph()

    # -- Program Outcomes ----------------------------------------------------
    if outcomes:
        doc.add_heading("Program Outcomes", 1)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        hdr[0].text = "Code"
        hdr[1].text = "Description"
        hdr[2].text = "Bloom Level"
        hdr[3].text = "Order"
        for o in outcomes:
            row = t.add_row()
            row.cells[0].text = o.code
            row.cells[1].text = o.description
            row.cells[2].text = o.bloom_level or "—"
            row.cells[3].text = str(o.display_order)
        doc.add_paragraph()

    # -- Course Structure (semester-grouped) ---------------------------------
    doc.add_heading("Course Structure", 1)
    by_semester = defaultdict(list)
    for c in courses:
        by_semester[c.semester].append(c)

    for sem in sorted(by_semester.keys()):
        doc.add_heading(f"Semester {sem}", 2)
        t = doc.add_table(rows=1, cols=6)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        hdr[0].text = "Code"
        hdr[1].text = "Title"
        hdr[2].text = "Credits"
        hdr[3].text = "Type"
        hdr[4].text = "L/T/P"
        hdr[5].text = "Prerequisites"
        for c in by_semester[sem]:
            ltp     = f"{c.hours_lecture or 0}/{c.hours_tutorial or 0}/{c.hours_practical or 0}"
            prereqs = ", ".join(prereqs_by_course.get(c.id, [])) or "—"
            row = t.add_row()
            row.cells[0].text = c.code
            row.cells[1].text = c.title
            row.cells[2].text = str(c.credits)
            row.cells[3].text = "Elective" if c.is_elective else "Core"
            row.cells[4].text = ltp
            row.cells[5].text = prereqs
        doc.add_paragraph()

    doc.save(buf)
